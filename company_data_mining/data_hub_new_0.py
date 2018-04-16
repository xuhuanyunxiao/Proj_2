# -*- coding: utf-8 -*-
"""
Created on Mon Mar 19 15:44:30 2018

@author: Administrator
"""
#%%
import os
import sys
import datetime as dt

today = dt.datetime.now().strftime('%Y%m%d')
names = locals()

import pandas as pd
import numpy as np

from sqlalchemy import create_engine
from pandas.io import sql

from impala.dbapi import connect
from impala.util import as_pandas

pd.set_option('precision', 3)
decimals = 3 # 小数位数

from docx import Document
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches

#%
pyfile_folder = r'D:\XH\Python_Project\Proj_2\files'
data_folder = r'D:\XH\Python_Project\Proj_2\data'
result_folder = r'D:\XH\Python_Project\Proj_2\result'

os.chdir(pyfile_folder)
sys.path.append(pyfile_folder)

from Tookits import specific_func  
from Tookits import cal_func

#% 建立连接
# MySQL
DB_CON_STR = 'mysql+pymysql://root:123456@localhost/sdm_2_mysql?charset=utf8'  
engine = create_engine(DB_CON_STR, echo=False) 

# Hive
conn = connect(host="192.168.20.102", port=10000,  # database="system", 
               auth_mechanism="PLAIN",
               user = 'admin', password = 'admin')
cursor = conn.cursor()

#%% 84张表（ES表结构）---找到字段解释
filename_list = pd.DataFrame(os.listdir(data_folder + '\\数据库表第23版-20180118'), columns = ['文件名称'])
filename_list['file_name'] = filename_list['文件名称'].apply(lambda x: os.path.splitext(x)[0])
filename_list['table_comment'] = filename_list['file_name'].apply(lambda x: x.split('-')[0])
filename_list['table_name'] = filename_list['file_name'].apply(lambda x: x.lower().split('-')[1])

index_list = []
restrain_list = []
ES_data = pd.DataFrame()

for index in filename_list.index:
    table_name = filename_list.loc[index, 'table_name']
    table_comment = filename_list.loc[index, 'table_comment']
    file_name = filename_list.loc[index, '文件名称']
    data_1 = pd.read_excel(os.path.join(data_folder + '\\数据库表第23版-20180118',
                                      file_name), header = None)
    list_filed = [str(x) for x in data_1.iloc[0,:].tolist()]
    
    if '数据库' in ''.join(list_filed):index_list.append(index)
    if index in index_list:
        header = 1
        field_n = '字段名称'
        field_t = '类型'
        field_c = '含义'
        field_e = '空'
        if index in [9,43]: field_e = '为空'
        if index == 30: field_c = '许可部门'
        if index == 40: 
            field_n = '名称'
            field_t = '数据类型'
            field_c = '注释'
            field_e = '非空'
        if index == 43: 
            field_n = '字段名'
            field_t = '字段类型'                
    else :
        header = 0
        field_n = '字段名'
        field_t = '字段类型'
        field_c = '含义'
        field_e = '能否为空'
        if index == 49:field_e = '为空'
    
    data = pd.read_excel(os.path.join(data_folder + '\\数据库表第23版-20180118',
                                      file_name), header = header) 
    data = data.drop_duplicates(subset=field_n, keep='last')
    data = pd.DataFrame(np.array(data),columns = data.columns,index = np.arange(data.shape[0]))
    data[field_e] = data[field_e].fillna('后来补充--')    
    data[field_t] = data[field_t].fillna('未定义')  
    data[field_t] = data[field_t].apply(lambda x:x.lower().strip().split('('))
    # VARCHAR2/varcahr--STRING;  number--NUMBER;  
    data['table_name'] = table_name
    ES_data = pd.concat([ES_data, data], axis = 0)        
    
    if index in index_list:
        data[field_e] = data[field_e].replace('√','否')
        restrain_name = data[data[field_e] == '否'][field_n].tolist()
    else :
        data[field_e] = data[field_e].replace('√','N')
        restrain_name = data[data[field_e] == 'N'][field_n].tolist()
    restrain_name.append('chanle_id')
    restrain_list.append(restrain_name)

ES_data = pd.DataFrame(np.array(ES_data),
                       columns = ES_data.columns,
                       index = np.arange(ES_data.shape[0]))
ES_data = ES_data.dropna(how = 'all', axis = 1)
ES_data = ES_data.drop('Unnamed: 11', axis = 1)

ES_data['field_name'] = ES_data['字段名称'].combine_first(ES_data['字段名'])
ES_data['field_name'] = ES_data['field_name'].combine_first(ES_data['名称'])

ES_data['field_type'] = ES_data['字段类型'].combine_first(ES_data['数据类型'])
ES_data['field_type'] = ES_data['field_type'].combine_first(ES_data['类型'])

ES_data['field_comment'] = ES_data['含义'].combine_first(ES_data['注释'])
ES_data['field_comment'] = ES_data['field_comment'].combine_first(ES_data['许可部门'])

ES_data['field_empty'] = ES_data['空'].combine_first(ES_data['为空'])
ES_data['field_empty'] = ES_data['field_empty'].combine_first(ES_data['能否为空'])
ES_data['field_empty'] = ES_data['field_empty'].combine_first(ES_data['非空'])

field_comment = ES_data[['field_name','field_comment']].drop_duplicates('field_name')
field_comment.columns = ['字段名', '字段解释'] 

#%% 53张表名
folder_name = 'DATA'
table_list_local = pd.DataFrame(os.listdir(data_folder + '\\' + folder_name), columns = ['文件名称'])
table_list_local['table_name'] = table_list_local['文件名称'].apply(lambda x: os.path.splitext(x)[0])

def run_hive_query(sql_command):   
    cursor.execute(sql_command)  
    return cursor.fetchall() 
database_name = 'data_hub_new'
cursor.execute("use "+ database_name) 
table_list_hive = [name[0] for name in run_hive_query("show tables")] 

filter_list = []
for name in table_list_local['table_name'].tolist():
    if name not in table_list_hive:
        print(name)
        filter_list.append(name)
        
#%
file_name = '17497公司名.txt'

file_path = os.path.join(data_folder,file_name)
encode = specific_func.get_txt_encode(file_path)
fid = open(file_path, encoding = encode)
company_num_local = fid.readlines()
fid.close()

company_num_local = pd.DataFrame(company_num_local, columns = ['company_name'])
company_num_local['company_name'] = company_num_local['company_name'].apply(lambda x: x.strip())
company_num_local['total_count'] = 1
company_num_local_unique = company_num_local.drop_duplicates('company_name')

table_name = 'company_base_business_merge_new'    
cursor.execute("select company_name,chanle_id from %s"%table_name)
company_num_hive = as_pandas(cursor)   
company_num_hive['total_count'] = 1
company_num_hive_unique = company_num_hive.drop_duplicates('company_name')

company_num_table = pd.DataFrame([['应有企业名单（local）',company_num_local.shape[0],
                                   company_num_local_unique.shape[0],
                                   company_num_local.shape[0] - company_num_local_unique.shape[0]],
                                  ['实有企业名单（hive）',company_num_hive.shape[0],
                                   company_num_hive_unique.shape[0],
                                   company_num_hive.shape[0] - company_num_hive_unique.shape[0]]],
                                  columns = ['企业数据', '原始数据', '去重后数据','相差'])
        
#%%
if not os.path.exists(result_folder + '\\' + folder_name + '\\' + today):
    os.makedirs(result_folder + '\\' + folder_name + '\\' + today) 
    
count_table = 0
table_stat_data = pd.DataFrame()
sum_table_data =  pd.DataFrame()
        
for table_name in table_list_local['table_name'].tolist():
    if table_name not in filter_list:
        print(table_name)
        cursor.execute("select * from %s"%table_name)
        names['%s' %table_name] = as_pandas(cursor) 
        names['%s' %table_name] = names['%s' %table_name].astype(str).replace('',np.nan)
        if 'rowkey' in  names['%s' %table_name].columns.tolist():
            names['%s' %table_name] = names['%s' %table_name].drop(['rowkey'], axis = 1)
        if names['%s' %table_name].iloc[0,0] == names['%s' %table_name].columns.tolist()[0]:
            names['%s' %table_name] = names['%s' %table_name].drop(0, axis = 0)
       
        # 统计
        fea_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, table_name + '.xlsx')        
        single_fea_desc = cal_func.describe(names['%s' %table_name],fea_filename, data_rate = 0.1)
        singlt_table = single_fea_desc[['是否缺失样本', '缺失量', '缺失率', 
                                        '现存量', '该特征含值的个数']]
        singlt_table = singlt_table.reset_index()
        singlt_table['table_name'] = table_name
        singlt_table['记录量'] = names['%s' %table_name].shape[0]
        singlt_table['缺失率'] = singlt_table['缺失率'].apply(lambda x: '%0.3f' % x)
        singlt_table.rename(columns={'index': '字段名'}, inplace=True) 
        singlt_table = pd.merge(singlt_table, field_comment, 
                                on = '字段名', how = 'left') 
            
        names['%s' %table_name]['single_count'] = 2  
        if 'company_name' not in names['%s' %table_name].columns.tolist():
            names['%s' %table_name] = pd.merge(company_num_hive[['company_name','chanle_id']], names['%s' %table_name],
                              on = 'chanle_id', how = 'right')             
        if table_name == 'company_branch_new':
            col = 'father_company_name'
            names['%s_unique' %table_name] = names['%s' %table_name].drop_duplicates(col)
            combined_local = pd.merge(company_num_local_unique, names['%s_unique' %table_name],
                                      left_on = 'company_name', 
                                      right_on = 'father_company_name', how = 'outer')        
            combined_local_same = combined_local[(combined_local['total_count'].notnull()) & 
                                       (combined_local['single_count'].notnull())]        
            combined_hive = pd.merge(company_num_hive_unique, names['%s_unique' %table_name],
                                     left_on = 'company_name', 
                                     right_on = 'father_company_name', how = 'outer')        
            combined_hive_same = combined_hive[(combined_hive['total_count'].notnull()) & 
                                       (combined_hive['single_count'].notnull())] 
        else :
            names['%s_unique' %table_name] = names['%s' %table_name].drop_duplicates('company_name')
            combined_local = pd.merge(company_num_local_unique, names['%s_unique' %table_name],
                              on = 'company_name', how = 'outer')        
            combined_local_same = combined_local[(combined_local['total_count'].notnull()) & 
                                       (combined_local['single_count'].notnull())]        
            combined_hive = pd.merge(company_num_hive_unique, names['%s_unique' %table_name],
                              on = 'company_name', how = 'outer')        
            combined_hive_same = combined_hive[(combined_hive['total_count'].notnull()) & 
                                       (combined_hive['single_count'].notnull())] 

        # 导出相同公司名
        file_name = result_folder + '\\' + folder_name + '\\' + today + '\\相同公司名单_%s.xlsx' %table_name
        with pd.ExcelWriter(file_name) as writer:
            combined_local_same.to_excel(writer,'local（一万七中）')
            combined_hive_same.to_excel(writer,'hive（一万四中）')
            writer.save()         
        
        table_stat_data = pd.concat([table_stat_data, singlt_table],axis = 0)
        
        missing_stat = singlt_table[singlt_table['是否缺失样本'] == 'YES']
        sum_table = pd.DataFrame([count_table + 1,table_name,
                                  names['%s' %table_name].shape[0],
                                  names['%s_unique' %table_name].shape[0],
                                  combined_local_same.shape[0],
                                  combined_hive_same.shape[0],
                                  names['%s' %table_name].shape[1] - 1,
                                  missing_stat['是否缺失样本'].count(),
                                  [missing_stat['缺失量'].min(),missing_stat['缺失量'].max()],
                                  [missing_stat['缺失率'].min(),
                                   missing_stat['缺失率'].max()]]).T
        sum_table.columns = ['序号','表名',
                             '原始数量','去重后数量',
                             '在应有企业名单中数量','在实有企业名单中数量',
                             '字段量','有缺失数据字段量',
                             '缺失量（min、max）',
                             '缺失率（min、max）']
        sum_table_data = pd.concat([sum_table_data, sum_table],axis = 0)
        
        del names['%s' %table_name]
        count_table += 1

#%% 初步统计
def add_table_doxc(data):
    table = document.add_table(rows=data.shape[0] + 1,cols=data.shape[1])
    for i in range(data.shape[0] + 1):
        hdr_cells=table.rows[i].cells
        for j in range(data.shape[1]):
            if i ==0:
                hdr_cells[j].text = data.columns.tolist()[j]
            else :
                hdr_cells[j].text = data.iloc[i-1,j]
    return hdr_cells
   
document = Document()
document.add_heading('52张表统计',0)

#增加表格： sum_table_data
document.add_heading('一、总体统计',1)
hdr_cells = add_table_doxc(company_num_table.astype(str))
document.add_paragraph('')
hdr_cells = add_table_doxc(sum_table_data.astype(str))

#增加表格： table_stat_data
document.add_heading('二、分表统计',1)
count_table = 1
for table_name in table_list_local['table_name'].tolist():
    sin_data = table_stat_data[table_stat_data['table_name'] == table_name]
    sin_data = sin_data[['字段名', '字段解释', '是否缺失样本', '缺失量', '缺失率', 
                         '现存量', '记录量', '该特征含值的个数']].astype(str)
    document.add_heading(str(count_table) + ' ' + table_name,2)
#    document.add_paragraph('样本总量为：%d'%sin_data.shape[0])
    document.add_paragraph('样本特征数为：%d'%sin_data.shape[0])
    hdr_cells = add_table_doxc(sin_data)
    count_table += 1
      
save_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, '52张表统计.docx')
document.save(os.path.join(save_filename))

table_stat_data['样本是否缺失'] = table_stat_data['是否缺失样本'].replace('YES',np.nan)
table_stat_data['样本是否缺失'] = table_stat_data['样本是否缺失'].combine_first(table_stat_data['缺失量'])

table_stat_data.to_excel(os.path.join(result_folder,'52张表统计.xlsx'))

#%%



