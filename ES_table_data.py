# -*- coding: utf-8 -*-
"""
Created on Tue Feb  6 13:36:06 2018

@author: Administrator
"""

#%%
import os
import sys
import time
import datetime as dt

today = dt.datetime.now().strftime('%Y%m%d')

import pandas as pd
import numpy as np

pd.set_option('precision', 3)
decimals = 3 # 小数位数

from docx import Document
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches

#%%
pyfile_folder = r'D:\XH\Python_Project\Proj_2\files'
data_folder = r'D:\XH\Python_Project\Proj_2\data'
result_folder = r'D:\XH\Python_Project\Proj_2\result'

os.chdir(pyfile_folder)
sys.path.append(pyfile_folder)

from Tookits import specific_func  
from Tookits import cal_func

#%% ES表结构 数据
folder_name = 'good_table'

filename_list = pd.DataFrame(os.listdir(data_folder + '\\' + folder_name), columns = ['文件名称'])
filename_list['file_name'] = filename_list['文件名称'].apply(lambda x: os.path.splitext(x)[0])

#%%
names = locals()
count = 1
except_list = []

if not os.path.exists(result_folder + '\\' + folder_name + '\\' + today):
    os.makedirs(result_folder + '\\' + folder_name + '\\' + today)  
    
num = count_table = 0
save_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, today + '_ES_HiveSQL_data.txt')
file = open(save_filename,"w")

table_stat_data = pd.DataFrame()
sum_table_data =  pd.DataFrame()

for index in filename_list.index:
    file_name = os.path.join(data_folder + '\\' + folder_name, filename_list['文件名称'][index])
    table_name = filename_list['file_name'][index]        
    encode = specific_func.get_txt_encode(file_name)
    data = pd.read_csv(file_name,nrows =1, encoding = encode)  #  
    
#        filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, today + '_' + table_name +'.csv')
#        data.to_csv(filename, sep='$', encoding='utf-8')
    
    print('==============================================================')
    num += 1 
    print('序号：',index)
    print('读取第 %d 个文件' % num)
    print(file_name)
    print(table_name)
    print('==============================================================')     
    
    file.write('USE data_hub_new;' + "\n")
    file.write('drop table if exists %s;' %table_name + "\n")
    file.write('CREATE EXTERNAL TABLE `%s` (' %table_name  + "\n")
    for ind, field_name in enumerate(data.columns.tolist()):
        if field_name =='number':
            field_name = field_name + '_NEW'
            print('************************************************')
            print(field_name)
            print(table_name)
            print('************************************************')
        field_type = 'STRING'        
        field_restrain = 'DEFAULT NULL'
        if ind == (data.shape[1]-1):
            punctuation = ')'
        else :
            punctuation = ','
        file.write('\t' * count + '%s' %field_name +
                   '\t' * count + '%s' % field_type +
                   '\t' * count + '%s' % field_restrain + punctuation + "\n")
#        file.write('PARTITIONED BY(dt STRING, country STRING)' + "\n")
    file.write("ROW FORMAT DELIMITED FIELDS TERMINATED BY ',' LOCATION '/tmp/20180208/%s';" %table_name + "\n")
    file.write("\n" * 2) 
               
    try :
        names['%s' %table_name] = pd.read_csv(file_name, encoding = encode) # , nrows =1
        
        if table_name == 'company_food_license':
            pass
        
        data_1 = pd.read_csv(file_name, encoding = encode)
        fea_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, table_name + '.xlsx')
        
        # 统计
        single_fea_desc = cal_func.describe(data_1,fea_filename, data_rate = 0.1)
        singlt_table = single_fea_desc[['是否缺失样本', '缺失量', '缺失率', 
                                        '现存量', '该特征含值的个数']]
        singlt_table = singlt_table.reset_index()
        singlt_table['table_name'] = table_name
        singlt_table['记录量'] = data_1.shape[0]
        singlt_table['缺失率'] = singlt_table['缺失率'].apply(lambda x: '%0.3f' % x)
        singlt_table.rename(columns={'index': '字段名'}, inplace=True) 
        
        table_stat_data = pd.concat([table_stat_data, singlt_table],axis = 0)
        
        missing_stat = singlt_table[singlt_table['是否缺失样本'] == 'YES']
        sum_table = pd.DataFrame([count_table + 1,table_name,data_1.shape[0],data_1.shape[1],
                                  missing_stat['是否缺失样本'].count(),
                                  [missing_stat['缺失量'].min(),missing_stat['缺失量'].max()],
                                  [missing_stat['缺失率'].min(),
                                   missing_stat['缺失率'].max()]]).T
        sum_table.columns = ['序号','表名','记录量','字段量','有缺失数据字段量',
                             '缺失量（min、max）','缺失率（min、max）']
        sum_table_data = pd.concat([sum_table_data, sum_table],axis = 0)
        
        print()
        print('*******  ', count_table + 1)
        count_table += 1
        
    except Exception as e:
        print('------------------------------')
        print(e)
        print(index)
        print(file_name)
        print(table_name)   
        except_list.append([index, file_name, table_name, e.args[0]])
        time.sleep(2)
        continue        
        
file.close()
filename_list['file_name'].to_excel(os.path.join(result_folder,'42_good_table_name.xlsx'))
save_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, 'a_出错文件及原因.xlsx')
pd.DataFrame(except_list,columns = ['index', 'file_path', 'table_name', 'Error']).to_excel(save_filename)

#%% 初步统计
def add_table_doxc(data):
    table = document.add_table(rows=data.shape[0] + 1,cols=data.shape[1])
    for i in range(data.shape[0] + 1):
        hdr_cells=table.rows[i].cells
        for j in range(data.shape[1]):
            if i ==0:
                hdr_cells[j].text = data.columns.tolist()[j]
            else :
                hdr_cells[j].text = data.iloc[i-1,j]
    return hdr_cells
   
document = Document()
document.add_heading('42张good_table（初步统计）',0)

#增加表格： sum_table_data
document.add_heading('一、总体统计',1)
hdr_cells = add_table_doxc(sum_table_data.astype(str))

#增加表格： table_stat_data
document.add_heading('二、分表统计',1)
for table_name,count_table in zip(sum_table_data['表名'].tolist(),sum_table_data['序号'].astype(str).tolist()):
    sin_data = table_stat_data[table_stat_data['table_name'] == table_name]
    sin_data = sin_data[['字段名','是否缺失样本', '缺失量', '缺失率', '现存量', '记录量','该特征含值的个数']].astype(str)
    document.add_heading(str(count_table) + ' ' + table_name,2)
#    document.add_paragraph('样本总量为：%d'%sin_data.shape[0])
    document.add_paragraph('样本特征数为：%d'%sin_data.shape[1])
    hdr_cells = add_table_doxc(sin_data)
      
save_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, '42张good_table（初步统计）.docx')
document.save(os.path.join(save_filename))

table_stat_data['样本是否缺失'] = table_stat_data['是否缺失样本'].replace('YES',np.nan)
table_stat_data['样本是否缺失'] = table_stat_data['样本是否缺失'].combine_first(table_stat_data['缺失量'])

table_stat_data.to_excel(os.path.join(result_folder,'42_good_table_stat.xlsx'))

#%%
save_filename = os.path.join(result_folder + '\\' + folder_name + '\\' + today, today + '_ES_put_data.txt')
file = open(save_filename,"w")

for index in filename_list.index:
    file_name = os.path.join(data_folder + '\\' + folder_name, filename_list['文件名称'][index])
    table_name = filename_list['file_name'][index]  
         
    file.write('hadoop fs -put /home/hadoop/Public/good_table/%s.csv /tmp/20180208/%s/%s.csv;' %(table_name,table_name,table_name))
    file.write("\n" * 2)                    
        
file.close()

#%%
#company_trademark_data = company_trademark.iloc[:,:-1].astype(str).applymap(lambda x:x.replace('\n','$$'))    
#company_trademark.iloc[:,:-1] =  company_trademark_data  

#company_trademark = company_trademark.astype(str).applymap(lambda x:x.replace('\n','$$'))    
#company_trademark.to_csv(os.path.join(r'D:\XH\Python_Project\Proj_2\result',
#                           'company_trademark.csv'), index=False, encoding='utf-8')
#
#company_solid_waste_imp = company_solid_waste_imp.astype(str).applymap(lambda x:x.replace('\r\n','$$'))    
#company_solid_waste_imp.to_csv(os.path.join(r'D:\XH\Python_Project\Proj_2\result',
#                           'company_solid_waste_imp.csv'), index=False, encoding='utf-8')
#
#
#
#company_base_business_merge_new.to_csv(os.path.join(r'D:\XH\Python_Project\Proj_2\result',
#                           'company_base_business_merge_new.csv'), index=False, encoding='utf-8')    
   #%% 
    
    
#c = company_base_business_merge_new['company_regis_code'][:5].apply(lambda x:'{:.0f}'.format(float(x)))    
    
    
    
    
    