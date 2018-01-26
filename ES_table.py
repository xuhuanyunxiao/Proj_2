# -*- coding: utf-8 -*-
"""
Created on Fri Jan 26 09:14:05 2018

@author: Administrator
"""

#%%
import os
import sys
import re
import time
import datetime as dt

today = dt.datetime.now().strftime('%Y%m%d')

import pandas as pd
import numpy as np
import xlrd

from docx import Document
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches

#%%
pyfile_folder = r'D:\XH\Python_Project\Proj_2\files'
data_folder = r'D:\XH\Python_Project\Proj_2\data'
result_folder = r'D:\XH\Python_Project\Proj_2\result'

os.chdir(pyfile_folder)
sys.path.append(pyfile_folder)

from Tookits import number_to_chinese  # 数字转中文
from Tookits import num_to_ROMAN_num  # 数字转罗马数字

#%% ES表结构 数据
file_path = os.path.join(data_folder, 'ES表结构.xlsx')
excel = xlrd.open_workbook(file_path)
sheet_names = [sheet.name for sheet in excel.sheets()]
  
#%% 打开文档
document = Document()

# 加入不同等级的标题
document.add_heading('原始数据（爬取（84）+外部（工商数据11、汇法网11））与ETL后数据（ES Table，32张表）',0)

document.add_heading('原始数据（爬取（84）+外部（工商数据11、汇法网11））',1)

# 外部数据采集表（参考）.xlsx
file_path = os.path.join(data_folder, '外部数据采集表（参考）.xlsx')
excel = xlrd.open_workbook(file_path)
sheet_names = [sheet.name for sheet in excel.sheets()]
for index, sheet_name in enumerate(sheet_names):  
    sheet_data = pd.read_excel(file_path, sheet_name).dropna(thresh=4, axis = 0)
    sheet_data['table_name'] = sheet_data['表中文名'] + '(' + sheet_data['表名'] + '）'
    sheet_data['head_name'] = sheet_data['中文别名'] + '  ---  ' + sheet_data['字段名']

    ind = num_to_ROMAN_num.transform_alabo2_roman_num(index + 1)    
    if sheet_name == '工商数据':
        document.add_heading(ind + ' ' + '外部（工商数据11）',2)
    else :
        document.add_heading(ind + ' ' + '外部（汇法网11）',2)
    
    for table_index, table_name in enumerate(sheet_data['table_name'].unique().tolist()):
        ind = number_to_chinese.num_to_chinese(table_index + 1)
        head_three = ind + '、' + table_name
        document.add_heading(head_three,3) 
        
        head_data = sheet_data[sheet_data['table_name'] == table_name]['head_name'].tolist()
        for head_index, head_name in enumerate(head_data):
            head_ind = str(head_index+1)
            head_four = head_ind + ' ' + head_name
            document.add_heading(head_four,4)     

# 数据库表第23版-20180118
filename_list = pd.DataFrame(os.listdir(data_folder + '\\数据库表第23版-20180118'), columns = ['文件名称'])
filename_list['文件名'] = filename_list['文件名称'].apply(lambda x: os.path.splitext(x)[0])
filename_list['扩展名'] = filename_list['文件名称'].apply(lambda x: os.path.splitext(x)[1])
            
file_path = os.path.join(data_folder, '数据索引分类.xlsx')
excel = xlrd.open_workbook(file_path)
sheet_names = [sheet.name for sheet in excel.sheets()]

for index, sheet_name in enumerate(sheet_names):  
    ind = num_to_ROMAN_num.transform_alabo2_roman_num(index + 3)    
    document.add_heading(ind + ' ' + sheet_name,2)
    if sheet_name == '其它':
        sheet_data = pd.read_excel(file_path, sheet_name, header = None)
        for file_index, filename in enumerate(sheet_data.iloc[:,0].tolist()):
            ind = number_to_chinese.num_to_chinese(file_index + 1)
            head_three = ind + '、' + table_name
            document.add_heading(head_three,3) 

            file_path = os.path.join(data_folder + '\\数据库表第23版-20180118', 
                                     list(filename_list[filename_list['文件名'] == filename]['文件名称'])[0])
            table_data = pd.read_excel(file_path)
    else :
        sheet_data = pd.read_excel(file_path, sheet_name, header = None)






# ES表结构.xlsx
file_path = os.path.join(data_folder, 'ES表结构.xlsx')
excel = xlrd.open_workbook(file_path)
sheet_names = [sheet.name for sheet in excel.sheets()]

document.add_heading('ETL后数据（ES Table，32张表）',1)
for index, sheet_name in enumerate(sheet_names):  
    sheet_data = pd.read_excel(file_path, sheet_name, skiprows=1).dropna(how = 'all', axis = 1).fillna('！！暂无！！')

    if sheet_name == '目录':
        ind = number_to_chinese.num_to_chinese(index+1)
        head_one = ind + '、' + sheet_name
        document.add_heading(head_one,2)    
        
        head_data = sheet_data['索引名称'].tolist()
        for head_index, head_name in enumerate(head_data):
            head_ind = str(head_index+2)
            head_two = head_ind + ' ' + head_name
            document.add_heading(head_two,3)   
    else :
        ind = number_to_chinese.num_to_chinese(index+1)
        head_one = ind + '、' + sheet_name +'（' + sheet_data.columns[1] + '）'
        document.add_heading(head_one,2)  
        
        head_data = sheet_data[sheet_data.columns[1]] + '  ---  ' + sheet_data[sheet_data.columns[2]]
        for head_index, head_name in enumerate(head_data.tolist()):
            if head_index > 0:
                head_ind = str(head_index)
                head_two = head_ind + ' ' + head_name
                document.add_heading(head_two,3)   

# 保存文件
filename = os.path.join(result_folder, today + '_' + 'ES_table.docx') 
document.save(filename)
#%%









